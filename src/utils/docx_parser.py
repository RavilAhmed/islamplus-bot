"""Парсер для извлечения уроков из DOCX файлов"""
import logging
from pathlib import Path
from typing import List, Dict, Optional
from docx import Document
from docx.shared import Inches
import re

logger = logging.getLogger(__name__)


def parse_docx_to_lessons(docx_path: Path) -> List[Dict]:
    """
    Парсит DOCX файл и извлекает уроки
    
    Ожидаемый формат:
    - Заголовки уровня 1 (Heading 1) = названия уроков
    - Текст под заголовком = содержание урока
    
    Args:
        docx_path: Путь к DOCX файлу
        
    Returns:
        Список словарей с уроками:
        [
            {
                "title": "Название урока",
                "content": "Текст урока",
                "day_number": 1
            },
            ...
        ]
    """
    try:
        doc = Document(docx_path)
        lessons = []
        current_lesson = None
        current_content = []
        day_number = 1
        
        for paragraph in doc.paragraphs:
            # Проверяем, является ли параграф заголовком уровня 1
            if paragraph.style.name.startswith('Heading 1') or (
                paragraph.style.name == 'Normal' and 
                paragraph.text.strip() and
                len(paragraph.text) < 100 and  # Короткий текст - вероятно заголовок
                paragraph.runs and
                paragraph.runs[0].bold  # Жирный текст - вероятно заголовок
            ):
                # Сохраняем предыдущий урок, если есть
                if current_lesson:
                    lessons.append({
                        "title": current_lesson,
                        "content": "\n\n".join(current_content).strip(),
                        "day_number": day_number
                    })
                    day_number += 1
                
                # Начинаем новый урок
                current_lesson = paragraph.text.strip()
                current_content = []
            else:
                # Добавляем текст к текущему уроку
                text = paragraph.text.strip()
                if text:
                    current_content.append(text)
        
        # Добавляем последний урок
        if current_lesson:
            lessons.append({
                "title": current_lesson,
                "content": "\n\n".join(current_content).strip(),
                "day_number": day_number
            })
        
        # Если не нашли уроки по заголовкам, разбиваем по разделам
        if not lessons:
            logger.warning("Не найдены заголовки, разбиваю по параграфам")
            all_text = "\n\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
            
            # Разбиваем на части по двойным переносам строк
            parts = re.split(r'\n\n+', all_text)
            
            for i, part in enumerate(parts, 1):
                if len(part.strip()) > 50:  # Игнорируем очень короткие части
                    # Первая строка = заголовок, остальное = содержание
                    lines = part.split('\n', 1)
                    title = lines[0].strip() if lines else f"Урок {i}"
                    content = lines[1].strip() if len(lines) > 1 else part.strip()
                    
                    lessons.append({
                        "title": title,
                        "content": content,
                        "day_number": i
                    })
        
        logger.info(f"Извлечено {len(lessons)} уроков из {docx_path}")
        return lessons
        
    except Exception as e:
        logger.error(f"Ошибка парсинга DOCX файла {docx_path}: {e}", exc_info=True)
        raise


def parse_docx_simple(docx_path: Path, split_by: str = "\n\n\n") -> List[Dict]:
    """
    Простой парсер: разбивает документ по разделителю
    
    Args:
        docx_path: Путь к DOCX файлу
        split_by: Разделитель между уроками
        
    Returns:
        Список уроков
    """
    try:
        doc = Document(docx_path)
        full_text = "\n".join([p.text for p in doc.paragraphs])
        
        # Разбиваем по разделителю
        parts = full_text.split(split_by)
        
        lessons = []
        for i, part in enumerate(parts, 1):
            part = part.strip()
            if not part:
                continue
            
            # Первая строка = заголовок
            lines = part.split('\n', 1)
            title = lines[0].strip() if lines else f"Урок {i}"
            content = lines[1].strip() if len(lines) > 1 else part.strip()
            
            lessons.append({
                "title": title,
                "content": content,
                "day_number": i
            })
        
        logger.info(f"Извлечено {len(lessons)} уроков (простой парсинг)")
        return lessons
        
    except Exception as e:
        logger.error(f"Ошибка простого парсинга DOCX: {e}", exc_info=True)
        raise
